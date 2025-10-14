"""
Work Order Checker Module

This module provides functionality to parse work order files and detect duplicate tasks.
Supports multiple file formats: TXT, CSV, JSON, PDF, HTML, XML, XLS, XLSX, DOC, DOCX
"""

import re
import json
import csv
from pathlib import Path
from typing import List, Dict, Set, Any
from dataclasses import dataclass

# Optional imports for different file formats
try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import xml.etree.ElementTree as ET
    HAS_XML = True
except ImportError:
    HAS_XML = False


@dataclass
class Task:
    """Represents a single task from a work order."""
    id: str
    description: str
    location: str
    raw_text: str
    
    def normalize(self) -> str:
        """Return a normalized version of the task for comparison."""
        # For duplicate detection, we focus on equipment ID and location
        # Extract equipment ID in brackets and location information
        equipment_match = re.search(r'\[([^\]]+)\]', self.raw_text)
        equipment_id = equipment_match.group(1) if equipment_match else ""
        
        # Extract location information (everything after the equipment ID)
        if equipment_match:
            location_part = self.raw_text[equipment_match.end():].strip()
        else:
            location_part = self.location or self.description
        
        # Combine equipment ID and location for comparison
        normalized = f"{equipment_id}|{location_part}".lower().strip()
        normalized = re.sub(r'\s+', ' ', normalized)
        
        return normalized
    
    def __hash__(self):
        return hash(self.normalize())
    
    def __eq__(self, other):
        if not isinstance(other, Task):
            return False
        return self.normalize() == other.normalize()


@dataclass
class WorkOrder:
    """Represents a work order containing multiple tasks."""
    name: str
    file_path: Path
    tasks: List[Task]
    
    def __str__(self):
        return f"{self.name} ({len(self.tasks)} tasks)"


class WorkOrderChecker:
    """Main class for checking work order duplicates."""
    
    def __init__(self):
        self.work_orders: List[WorkOrder] = []
        self.all_tasks: Dict[Task, List[str]] = {}  # Task -> List of work order names
    
    def load_work_order(self, file_path: Path) -> WorkOrder:
        """Load a work order from a file."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type and parse accordingly
        suffix = file_path.suffix.lower()
        
        if suffix == '.json':
            work_order = self._parse_json_work_order(file_path)
        elif suffix == '.csv':
            work_order = self._parse_csv_work_order(file_path)
        elif suffix in ['.xls', '.xlsx']:
            work_order = self._parse_excel_work_order(file_path)
        elif suffix == '.pdf':
            work_order = self._parse_pdf_work_order(file_path)
        elif suffix in ['.html', '.htm', '.HTM']:
            work_order = self._parse_html_work_order(file_path)
        elif suffix == '.xml':
            work_order = self._parse_xml_work_order(file_path)
        elif suffix in ['.doc', '.docx']:
            work_order = self._parse_doc_work_order(file_path)
        else:
            # Default to text parsing
            work_order = self._parse_text_work_order(file_path)
        
        self.work_orders.append(work_order)
        
        # Index tasks for duplicate detection
        for task in work_order.tasks:
            if task not in self.all_tasks:
                self.all_tasks[task] = []
            self.all_tasks[task].append(work_order.name)
        
        return work_order
    
    def _parse_text_work_order(self, file_path: Path) -> WorkOrder:
        """Parse a text-based work order file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        tasks = []
        
        # Look for patterns like "exit lights [212934] MOB B - ground floor - HRC back door"
        # Focus on equipment with location information
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Look for equipment patterns with IDs and locations
            equipment_pattern = r'(exit lights?|emergency lighting?|fire extinguisher)\s*\[([^\]]+)\]\s*(.+)'
            
            match = re.match(equipment_pattern, line, re.IGNORECASE)
            if match:
                equipment_type = match.group(1).strip()
                equipment_id = match.group(2).strip()
                location = match.group(3).strip()
                
                task = Task(
                    id=equipment_id,
                    description=equipment_type,
                    location=location,
                    raw_text=line
                )
                tasks.append(task)
            else:
                # Fall back to general pattern matching
                general_pattern = r'(.+?)\[([^\]]+)\](.+)'
                match = re.match(general_pattern, line)
                if match:
                    description = match.group(1).strip()
                    task_id = match.group(2).strip()
                    location = match.group(3).strip()
                    
                    task = Task(
                        id=task_id,
                        description=description,
                        location=location,
                        raw_text=line
                    )
                    tasks.append(task)
        
        return WorkOrder(
            name=file_path.stem,
            file_path=file_path,
            tasks=tasks
        )
    
    def _parse_csv_work_order(self, file_path: Path) -> WorkOrder:
        """Parse a CSV work order file."""
        tasks = []
        
        with open(file_path, 'r', encoding='utf-8') as f:
            # Try to detect if first row is header
            sample = f.read(1024)
            f.seek(0)
            
            reader = csv.reader(f)
            rows = list(reader)
            
            if not rows:
                return WorkOrder(file_path.stem, file_path, [])
            
            # Assume first row might be header if it contains common header words
            header_indicators = ['id', 'task', 'description', 'location', 'item']
            first_row_lower = [cell.lower() for cell in rows[0]]
            has_header = any(indicator in ' '.join(first_row_lower) for indicator in header_indicators)
            
            start_row = 1 if has_header else 0
            
            for i, row in enumerate(rows[start_row:], start_row):
                if not row or not any(cell.strip() for cell in row):
                    continue
                
                # Join all non-empty cells to form the task description
                task_text = ' '.join(cell.strip() for cell in row if cell.strip())
                
                # Extract ID if possible
                numbers = re.findall(r'\d+', task_text)
                task_id = numbers[0] if numbers else f"row_{i}"
                
                task = Task(
                    id=task_id,
                    description=task_text,
                    location="",
                    raw_text=task_text
                )
                tasks.append(task)
        
        return WorkOrder(
            name=file_path.stem,
            file_path=file_path,
            tasks=tasks
        )
    
    def _parse_json_work_order(self, file_path: Path) -> WorkOrder:
        """Parse a JSON work order file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        tasks = []
        
        # Handle different JSON structures
        if isinstance(data, list):
            # Array of tasks
            for item in data:
                task = self._json_item_to_task(item)
                if task:
                    tasks.append(task)
        elif isinstance(data, dict):
            # Object with tasks array
            if 'tasks' in data and isinstance(data['tasks'], list):
                for item in data['tasks']:
                    task = self._json_item_to_task(item)
                    if task:
                        tasks.append(task)
            else:
                # Single task object
                task = self._json_item_to_task(data)
                if task:
                    tasks.append(task)
        
        return WorkOrder(
            name=file_path.stem,
            file_path=file_path,
            tasks=tasks
        )
    
    def _json_item_to_task(self, item: Any) -> Task:
        """Convert a JSON item to a Task object."""
        if isinstance(item, str):
            # Simple string task
            numbers = re.findall(r'\d+', item)
            return Task(
                id=numbers[0] if numbers else "unknown",
                description=item,
                location="",
                raw_text=item
            )
        elif isinstance(item, dict):
            # Task object
            task_id = str(item.get('id', item.get('task_id', 'unknown')))
            description = str(item.get('description', item.get('task', item.get('name', ''))))
            location = str(item.get('location', item.get('area', '')))
            
            # If no description found, join all values
            if not description:
                description = ' '.join(str(v) for v in item.values() if v)
            
            return Task(
                id=task_id,
                description=description,
                location=location,
                raw_text=description
            )
        
        return None
    
    def find_duplicates(self) -> List[Dict[str, Any]]:
        """Find all duplicate tasks across work orders."""
        duplicates = []
        
        for task, work_order_names in self.all_tasks.items():
            if len(work_order_names) > 1:
                duplicates.append({
                    'task': task.raw_text,
                    'task_id': task.id,
                    'normalized': task.normalize(),
                    'work_orders': work_order_names,
                    'count': len(work_order_names)
                })
        
        # Sort by number of duplicates (most duplicated first)
        duplicates.sort(key=lambda x: x['count'], reverse=True)
        
        return duplicates
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get statistics about the loaded work orders."""
        total_tasks = sum(len(wo.tasks) for wo in self.work_orders)
        unique_tasks = len(self.all_tasks)
        duplicate_tasks = sum(1 for tasks in self.all_tasks.values() if len(tasks) > 1)
        
        return {
            'work_orders_count': len(self.work_orders),
            'total_tasks': total_tasks,
            'unique_tasks': unique_tasks,
            'duplicate_tasks': duplicate_tasks,
            'duplication_rate': (duplicate_tasks / unique_tasks * 100) if unique_tasks > 0 else 0
        }
    
    def export_duplicates_report(self, output_path: Path) -> None:
        """Export a detailed duplicates report to a file."""
        duplicates = self.find_duplicates()
        stats = self.get_statistics()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("Work Order Duplicate Analysis Report\n")
            f.write("=" * 50 + "\n\n")
            
            f.write("Summary:\n")
            f.write(f"- Total Work Orders: {stats['work_orders_count']}\n")
            f.write(f"- Total Tasks: {stats['total_tasks']}\n")
            f.write(f"- Unique Tasks: {stats['unique_tasks']}\n")
            f.write(f"- Duplicate Tasks: {stats['duplicate_tasks']}\n")
            f.write(f"- Duplication Rate: {stats['duplication_rate']:.1f}%\n\n")
            
            if duplicates:
                f.write(f"Duplicate Tasks Found ({len(duplicates)}):\n")
                f.write("-" * 50 + "\n")
                
                for i, duplicate in enumerate(duplicates, 1):
                    f.write(f"\n{i}. Task ID: {duplicate['task_id']}\n")
                    f.write(f"   Description: {duplicate['task']}\n")
                    f.write(f"   Appears in {duplicate['count']} work orders:\n")
                    for wo_name in duplicate['work_orders']:
                        f.write(f"     - {wo_name}\n")
            else:
                f.write("No duplicate tasks found!\n")
    
    def _parse_excel_work_order(self, file_path: Path) -> WorkOrder:
        """Parse an Excel work order file (.xls or .xlsx)."""
        if not HAS_PANDAS:
            raise ImportError("pandas is required for Excel file support. Install with: pip install pandas openpyxl xlrd")
        
        try:
            # Try to read Excel file
            df = pd.read_excel(file_path, engine='openpyxl' if file_path.suffix == '.xlsx' else 'xlrd')
            tasks = []
            
            for index, row in df.iterrows():
                # Convert row to string, joining non-null values
                row_values = [str(val) for val in row.values if pd.notna(val)]
                if not row_values:
                    continue
                
                task_text = ' '.join(row_values)
                
                # Extract ID if possible
                numbers = re.findall(r'\d+', task_text)
                task_id = numbers[0] if numbers else f"row_{index}"
                
                task = Task(
                    id=task_id,
                    description=task_text,
                    location="",
                    raw_text=task_text
                )
                tasks.append(task)
            
            return WorkOrder(
                name=file_path.stem,
                file_path=file_path,
                tasks=tasks
            )
        except Exception as e:
            raise Exception(f"Error parsing Excel file {file_path}: {e}")
    
    def _parse_pdf_work_order(self, file_path: Path) -> WorkOrder:
        """Parse a PDF work order file."""
        if not HAS_PDF:
            raise ImportError("PyPDF2 is required for PDF file support. Install with: pip install PyPDF2")
        
        try:
            tasks = []
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                # Extract text from all pages
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                # Parse the extracted text using the same logic as text files
                lines = text_content.split('\n')
                task_pattern = r'(.+?)\[(\d+)\](.+)'
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Try to match the pattern with ID in brackets
                    match = re.match(task_pattern, line)
                    if match:
                        description = match.group(1).strip()
                        task_id = match.group(2)
                        location = match.group(3).strip()
                        
                        task = Task(
                            id=task_id,
                            description=description,
                            location=location,
                            raw_text=line
                        )
                        tasks.append(task)
                    else:
                        # If no pattern match, check if line looks like a task
                        numbers = re.findall(r'\d+', line)
                        if numbers and len(line) > 10:  # Reasonable task length
                            task_id = numbers[0]
                            task = Task(
                                id=task_id,
                                description=line,
                                location="",
                                raw_text=line
                            )
                            tasks.append(task)
            
            return WorkOrder(
                name=file_path.stem,
                file_path=file_path,
                tasks=tasks
            )
        except Exception as e:
            raise Exception(f"Error parsing PDF file {file_path}: {e}")
    
    def _parse_html_work_order(self, file_path: Path) -> WorkOrder:
        """Parse an HTML work order file."""
        if not HAS_BS4:
            raise ImportError("beautifulsoup4 is required for HTML file support. Install with: pip install beautifulsoup4")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                
                tasks = []
                
                # Look specifically for equipment entries in table cells
                # Pattern: <b>Exit Light [ID]</b> Location description
                table_cells = soup.find_all('td', class_='data_underline')
                
                for cell in table_cells:
                    cell_text = cell.get_text(separator=' ', strip=True)
                    
                    # Look for equipment with bold formatting and IDs
                    # Pattern: Exit Light [ID] Location: Building details
                    equipment_pattern = r'(Exit Light|Fire Extinguisher|Emergency Light)(?:\s+[\w-]+)?\s*\[([^\]]+)\]\s*(.+?)(?:\s*MAIN HOSPITAL|\s*Equipment Lists|\s*Exit Lights|$)'
                    
                    match = re.search(equipment_pattern, cell_text, re.IGNORECASE)
                    if match:
                        equipment_type = match.group(1).strip()
                        equipment_id = match.group(2).strip()
                        location = match.group(3).strip()
                        
                        # Clean up location - remove trailing colons and extra building info
                        location = re.sub(r':\s*$', '', location)
                        location = re.sub(r'\s+', ' ', location).strip()
                        
                        if location:  # Only add if we have a meaningful location
                            task_text = f"{equipment_type} [{equipment_id}] {location}"
                            
                            task = Task(
                                id=equipment_id,
                                description=equipment_type,
                                location=location,
                                raw_text=task_text
                            )
                            tasks.append(task)
                
                return WorkOrder(
                    name=file_path.stem,
                    file_path=file_path,
                    tasks=tasks
                )
        except Exception as e:
            raise Exception(f"Error parsing HTML file {file_path}: {e}")
    
    def _parse_xml_work_order(self, file_path: Path) -> WorkOrder:
        """Parse an XML work order file."""
        if not HAS_XML:
            raise ImportError("xml.etree.ElementTree is required for XML file support (usually included with Python)")
        
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            tasks = []
            
            # Look for common XML structures
            # Try to find elements that might contain task information
            for element in root.iter():
                if element.text and element.text.strip():
                    text = element.text.strip()
                    
                    # Skip very short text or common XML tags
                    if len(text) < 5 or text.lower() in ['true', 'false', 'yes', 'no']:
                        continue
                    
                    # Look for task patterns
                    task_pattern = r'(.+?)\[(\d+)\](.+)'
                    match = re.match(task_pattern, text)
                    
                    if match:
                        description = match.group(1).strip()
                        task_id = match.group(2)
                        location = match.group(3).strip()
                        
                        task = Task(
                            id=task_id,
                            description=description,
                            location=location,
                            raw_text=text
                        )
                        tasks.append(task)
                    else:
                        # Check for potential task IDs
                        numbers = re.findall(r'\d+', text)
                        if numbers and len(text) > 10:
                            task_id = numbers[0]
                            task = Task(
                                id=task_id,
                                description=text,
                                location="",
                                raw_text=text
                            )
                            tasks.append(task)
            
            return WorkOrder(
                name=file_path.stem,
                file_path=file_path,
                tasks=tasks
            )
        except Exception as e:
            raise Exception(f"Error parsing XML file {file_path}: {e}")
    
    def _parse_doc_work_order(self, file_path: Path) -> WorkOrder:
        """Parse a Word document work order file (.doc or .docx)."""
        if not HAS_DOCX:
            raise ImportError("python-docx is required for Word document support. Install with: pip install python-docx")
        
        if file_path.suffix.lower() == '.doc':
            raise Exception("Legacy .doc files are not supported. Please convert to .docx format.")
        
        try:
            doc = Document(file_path)
            tasks = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text or len(text) < 5:
                    continue
                
                # Look for task patterns
                task_pattern = r'(.+?)\[(\d+)\](.+)'
                match = re.match(task_pattern, text)
                
                if match:
                    description = match.group(1).strip()
                    task_id = match.group(2)
                    location = match.group(3).strip()
                    
                    task = Task(
                        id=task_id,
                        description=description,
                        location=location,
                        raw_text=text
                    )
                    tasks.append(task)
                else:
                    # Check for potential tasks with numbers
                    numbers = re.findall(r'\d+', text)
                    if numbers and len(text) > 10:
                        task_id = numbers[0]
                        task = Task(
                            id=task_id,
                            description=text,
                            location="",
                            raw_text=text
                        )
                        tasks.append(task)
            
            # Also check tables if they exist
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if len(row_text) > 10:
                        numbers = re.findall(r'\d+', row_text)
                        if numbers:
                            task_id = numbers[0]
                            task = Task(
                                id=task_id,
                                description=row_text,
                                location="",
                                raw_text=row_text
                            )
                            tasks.append(task)
            
            return WorkOrder(
                name=file_path.stem,
                file_path=file_path,
                tasks=tasks
            )
        except Exception as e:
            raise Exception(f"Error parsing Word document {file_path}: {e}")